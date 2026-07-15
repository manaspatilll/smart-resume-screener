import io
import pdfplumber
import docx
from fastapi import UploadFile, HTTPException


def extract_text_from_upload(file: UploadFile, content: bytes) -> str:
    """Dispatch to the right extractor based on file extension."""
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        return _extract_pdf(content)
    elif filename.endswith(".docx"):
        return _extract_docx(content)
    elif filename.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {filename}. Use PDF, DOCX, or TXT.",
        )


def _extract_pdf(content: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    text = "\n".join(text_parts).strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="Could not extract text from PDF (it may be a scanned image).",
        )
    return text


def _extract_docx(content: bytes) -> str:
    document = docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    text = "\n".join(paragraphs).strip()
    if not text:
        raise HTTPException(status_code=422, detail="Could not extract text from DOCX.")
    return text
